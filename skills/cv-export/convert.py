#!/usr/bin/env python3
"""
cv-export/convert.py — Markdown CV → template-matched .docx.

Uses template.docx as the formatting source of truth. For each element type
(name, section header, company line, bullet, etc.) the paragraph properties
(spacing, indents, tab stops, numPr) are deep-copied from the corresponding
prototype paragraph in the template, then fresh runs are added with the correct
text and inline formatting. This means any change to the template automatically
propagates — no hardcoded spacing values in this file.

Template must exist at <input-dir>/template.docx.

Usage:
    python3 convert.py <input.md> [<output.docx>]

Defaults output to <input-dir>/build/<input-stem>.docx.

Markdown mapping:
  # H1              → name (para 0)
  ## H2             → section header (para 4)
  ### H3            → company/role block; following *italic* line → date/location
  **bold line**     → project sub-header (exp: para 19) or project card (projects: para 46)
  **X** — *Y*       → education institution (para 62)
  - bullet          → bullet (para 20), skill line (para 8), or edu achievement (para 64)
  plain text        → body (para 5) or subtitle/contact/links in header
  ---               → ignored
"""

from __future__ import annotations

import copy
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ── Colours ───────────────────────────────────────────────────────────────
C_DARK  = RGBColor(0x1A, 0x1A, 0x2E)
C_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
C_MUTED = RGBColor(0x77, 0x77, 0x77)
C_SEP   = RGBColor(0xBB, 0xBB, 0xBB)
C_LINK  = RGBColor(0x05, 0x63, 0xC1)

# ── Sizes ─────────────────────────────────────────────────────────────────
SZ_NAME    = Pt(17)
SZ_SECTION = Pt(10.5)
SZ_COMPANY = Pt(10.5)
SZ_BODY    = Pt(9.5)
SZ_CONTACT = Pt(9)
FONT = "Calibri"

# ── Prototype paragraph indices in template.docx ───────────────────────
#   Para  0 — name
#   Para  1 — subtitle (role tagline below name)
#   Para  2 — contact line  (Belfast · email)
#   Para  3 — links line    (github · linkedin)
#   Para  4 — section header
#   Para  5 — body paragraph
#   Para  8 — skill line    (bold label: content)
#   Para 17 — company line  (bold name \t date)
#   Para 18 — role line     (italic role \t location)
#   Para 19 — exp project header
#   Para 20 — bullet
#   Para 46 — personal project header (bold name — italic tagline)
#   Para 62 — education institution
#   Para 63 — education degree
#   Para 64 — education achievement
PROTO_IDX = {
    "name":        0,
    "subtitle":    1,
    "contact":     2,
    "links":       3,
    "section":     4,
    "body":        5,
    "skill":       8,
    "company":     17,
    "role":        18,
    "project_exp": 19,
    "bullet":      20,
    "project_pers":46,
    "edu_inst":    62,
    "edu_degree":  63,
    "edu_achieve": 64,
}

_PPR: dict[str, object] = {}   # keyed by proto name → deep-copied pPr element


def load_prototypes(template_path: Path) -> None:
    doc = Document(str(template_path))
    paras = doc.paragraphs
    for name, idx in PROTO_IDX.items():
        if idx < len(paras):
            pPr = paras[idx]._p.find(qn("w:pPr"))
            _PPR[name] = copy.deepcopy(pPr) if pPr is not None else None
        else:
            _PPR[name] = None


def mkp(doc: Document, proto: str):
    """New paragraph with pPr deep-copied from the named prototype."""
    p = doc.add_paragraph()
    src = _PPR.get(proto)
    if src is not None:
        new_pPr = copy.deepcopy(src)
        dst = p._p.find(qn("w:pPr"))
        if dst is not None:
            p._p.replace(dst, new_pPr)
        else:
            p._p.insert(0, new_pPr)
    return p


# ── Inline parser ─────────────────────────────────────────────────────────
def parse_inline(text: str) -> list:
    tokens: list = []
    buf = ""
    i = 0

    def flush():
        nonlocal buf
        if buf:
            tokens.append(("text", buf))
            buf = ""

    while i < len(text):
        c = text[i]
        if c == "`":
            end = text.find("`", i + 1)
            if end > -1:
                flush()
                tokens.append(("code", text[i + 1:end]))
                i = end + 1
                continue
        if c == "[":
            close = text.find("]", i)
            if close > -1 and close + 1 < len(text) and text[close + 1] == "(":
                url_end = text.find(")", close)
                if url_end > -1:
                    flush()
                    tokens.append(("link", text[i + 1:close], text[close + 2:url_end]))
                    i = url_end + 1
                    continue
        if text[i:i + 2] == "**":
            end = text.find("**", i + 2)
            if end > -1:
                flush()
                tokens.append(("bold", text[i + 2:end]))
                i = end + 2
                continue
        if c == "*":
            end = text.find("*", i + 1)
            if end > -1:
                flush()
                tokens.append(("italic", text[i + 1:end]))
                i = end + 1
                continue
        buf += c
        i += 1
    flush()
    return tokens


# ── Run / element helpers ─────────────────────────────────────────────────
def mkrun(p, text, bold=False, italic=False, size=None, color=None):
    run = p.add_run(text)
    run.font.name = FONT
    if size:
        run.font.size = size
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    return run


def add_hyperlink(p, text, url, size=SZ_CONTACT):
    part = p.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    wr = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size.pt * 2)))
    rPr.append(sz)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), str(C_LINK))
    rPr.append(color_el)
    wr.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    wr.append(t)
    hl.append(wr)
    p._p.append(hl)


def add_inline(p, text, size=SZ_BODY, default_color=C_DARK):
    for tok in parse_inline(text):
        kind = tok[0]
        if kind == "text":
            mkrun(p, tok[1], size=size, color=default_color)
        elif kind == "bold":
            mkrun(p, tok[1], bold=True, size=size, color=default_color)
        elif kind == "italic":
            mkrun(p, tok[1], italic=True, size=size, color=C_MUTED)
        elif kind == "code":
            r = p.add_run(tok[1])
            r.font.name = "Consolas"
            r.font.size = size
        elif kind == "link":
            add_hyperlink(p, tok[1], tok[2], size=size)


# ── Paragraph builders ────────────────────────────────────────────────────

def add_name(doc, text):
    p = mkp(doc, "name")
    mkrun(p, text, bold=True, size=SZ_NAME, color=C_DARK)


def add_subtitle(doc, text):
    clean = re.sub(r"^\*\*|\*\*$", "", text.strip())
    p = mkp(doc, "subtitle")
    mkrun(p, clean, size=SZ_BODY, color=C_MUTED)


def add_contact(doc, text):
    p = mkp(doc, "contact")
    parts = re.split(r"\s*·\s*", text)
    for i, part in enumerate(parts):
        if i > 0:
            mkrun(p, " · ", size=SZ_CONTACT, color=C_SEP)
        mkrun(p, part.strip(), size=SZ_CONTACT, color=C_MUTED)


def add_links(doc, text):
    p = mkp(doc, "links")
    parts = re.split(r"\s*·\s*", text)
    for i, part in enumerate(parts):
        if i > 0:
            mkrun(p, "  ·  ", size=SZ_CONTACT, color=C_SEP)
        part = part.strip()
        m = re.match(r"^\[(.+?)\]\((.+?)\)$", part)
        if m:
            add_hyperlink(p, m.group(1), m.group(2), size=SZ_CONTACT)
        else:
            add_inline(p, part, size=SZ_CONTACT, default_color=C_MUTED)


def add_section(doc, text):
    p = mkp(doc, "section")
    mkrun(p, text.upper(), bold=True, size=SZ_SECTION, color=C_BLUE)


def add_company(doc, company, date):
    p = mkp(doc, "company")
    mkrun(p, company, bold=True, size=SZ_COMPANY, color=C_DARK)
    if date:
        mkrun(p, "\t" + date, size=SZ_BODY, color=C_MUTED)


def add_role(doc, role, location):
    p = mkp(doc, "role")
    text = role + ("\t" + location if location else "")
    mkrun(p, text, italic=True, size=SZ_BODY, color=C_MUTED)


def add_project_exp(doc, text):
    p = mkp(doc, "project_exp")
    add_inline(p, text, size=SZ_BODY, default_color=C_DARK)
    for run in p.runs:
        run.bold = True


def add_project_personal(doc, text):
    clean = re.sub(r"^\*\*|\*\*$", "", text.strip())
    p = mkp(doc, "project_pers")
    if " — " in clean:
        name, tag = clean.split(" — ", 1)
        mkrun(p, name, bold=True, size=SZ_BODY, color=C_DARK)
        mkrun(p, " — " + tag, italic=True, size=SZ_BODY, color=C_MUTED)
    else:
        mkrun(p, clean, bold=True, size=SZ_BODY, color=C_DARK)


def add_edu_inst(doc, bold_part, location):
    p = mkp(doc, "edu_inst")
    mkrun(p, bold_part, bold=True, size=SZ_BODY, color=C_DARK)
    if location:
        mkrun(p, "\t" + location, size=SZ_BODY, color=C_MUTED)


def add_edu_degree(doc, text):
    p = mkp(doc, "edu_degree")
    mkrun(p, text, italic=True, size=SZ_BODY, color=C_MUTED)


def add_bullet(doc, text):
    p = mkp(doc, "bullet")
    add_inline(p, text, size=SZ_BODY, default_color=C_DARK)


def add_skill_line(doc, text):
    p = mkp(doc, "skill")
    add_inline(p, text, size=SZ_BODY, default_color=C_DARK)


def add_edu_achieve(doc, text):
    p = mkp(doc, "edu_achieve")
    add_inline(p, text, size=SZ_BODY, default_color=C_DARK)


def add_body(doc, text):
    p = mkp(doc, "body")
    add_inline(p, text, size=SZ_BODY, default_color=C_DARK)


# ── Document helpers ──────────────────────────────────────────────────────

def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


# ── Regexes ───────────────────────────────────────────────────────────────
HEADING_RE  = re.compile(r"^(#{1,6})\s+(.*)")
BULLET_RE   = re.compile(r"^[-*+]\s+(.*)")
ITALIC_LINE = re.compile(r"^\*(.+)\*$")
BOLD_LINE   = re.compile(r"^\*\*([^*]+)\*\*$")
EDU_INST_RE = re.compile(r"^\*\*(.+?)\*\*\s*—\s*\*(.+?)\*$")

SECTION_MAP = {
    "PROFESSIONAL SUMMARY": "summary",
    "TECHNICAL SKILLS": "skills",
    "PROFESSIONAL EXPERIENCE": "experience",
    "PERSONAL PROJECT WORK": "projects",
    "EDUCATION & ACHIEVEMENTS": "education",
}


# ── Conversion ────────────────────────────────────────────────────────────

def convert(md_path: Path, docx_path: Path, template_path: Path | None = None) -> None:
    if template_path is None:
        template_path = md_path.parent / "template.docx"

    if not template_path.exists():
        print(f"Error: template not found: {template_path}", file=sys.stderr)
        print("Place template.docx alongside the input .md file.", file=sys.stderr)
        sys.exit(1)

    load_prototypes(template_path)

    doc = Document(str(template_path))
    clear_body(doc)

    lines = md_path.read_text(encoding="utf-8").split("\n")
    section_ctx = "header"
    header_state = 0   # 0=name  1=subtitle  2=contact  3=links
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        i += 1

        if not line or line in ("---", "***"):
            continue

        h = HEADING_RE.match(line)
        if h:
            level, text = len(h.group(1)), h.group(2).strip()
            if level == 1:
                add_name(doc, text)
                header_state = 1
            elif level == 2:
                section_ctx = SECTION_MAP.get(text.upper(), "other")
                add_section(doc, text)
            elif level == 3:
                company, role = (text.split(" — ", 1) if " — " in text else (text, ""))
                date = location = ""
                if i < len(lines):
                    nxt = lines[i].strip()
                    m = ITALIC_LINE.match(nxt)
                    if m:
                        parts = re.split(r"\s*·\s*", m.group(1))
                        location = parts[0]
                        date = parts[1] if len(parts) >= 2 else ""
                        i += 1
                add_company(doc, company, date)
                if role:
                    add_role(doc, role, location)
            continue

        b = BULLET_RE.match(line)
        if b:
            content = b.group(1)
            if section_ctx == "skills":
                add_skill_line(doc, content)
            elif section_ctx == "education":
                add_edu_achieve(doc, content)
            else:
                add_bullet(doc, content)
            continue

        # Non-heading, non-bullet
        if header_state == 1:
            add_subtitle(doc, line)
            header_state = 2
        elif header_state == 2:
            add_contact(doc, line)
            header_state = 3
        elif header_state == 3:
            add_links(doc, line)
            header_state = 4
        elif section_ctx == "education":
            m = EDU_INST_RE.match(line)
            if m:
                add_edu_inst(doc, m.group(1), m.group(2))
            elif BOLD_LINE.match(line):
                add_project_exp(doc, line)
            else:
                add_edu_degree(doc, line)
        elif section_ctx in ("experience", "projects"):
            if BOLD_LINE.match(line):
                if section_ctx == "projects":
                    add_project_personal(doc, line)
                else:
                    add_project_exp(doc, line)
            else:
                add_body(doc, line)
        else:
            add_body(doc, line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main():
    args = sys.argv[1:]
    if not args:
        print("Usage: convert.py <input.md> [<output.docx>]", file=sys.stderr)
        sys.exit(1)

    md_path = Path(args[0]).expanduser().resolve()
    if not md_path.exists():
        print(f"Error: input not found: {md_path}", file=sys.stderr)
        sys.exit(1)

    docx_path = (
        Path(args[1]).expanduser().resolve()
        if len(args) >= 2
        else md_path.parent / "build" / (md_path.stem + ".docx")
    )

    convert(md_path, docx_path)
    print(str(docx_path))


if __name__ == "__main__":
    main()
